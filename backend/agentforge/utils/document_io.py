"""Read and write PDF and Word documents as plain text."""

from __future__ import annotations

from pathlib import Path

from agentforge.utils.optional_deps import OptionalDependencyError, ensure_document_packages

DOCUMENT_SUFFIXES = frozenset({".pdf", ".docx"})


def is_document_path(path: str | Path) -> bool:
    """
    Return True when the path uses a supported document extension.

    :param path: File path or suffix-bearing name
    :return: Whether document I/O helpers apply
    """
    return Path(path).suffix.lower() in DOCUMENT_SUFFIXES


def read_document_text(path: Path, *, auto_install: bool = True) -> str:
    """
    Extract plain text from a PDF or DOCX file.

    :param path: Absolute path to the document
    :param auto_install: Attempt pip install when dependencies are missing
    :return: Extracted text content
    :raises OptionalDependencyError: When dependencies are unavailable
    :raises ValueError: When the file type is unsupported
    """
    suffix = path.suffix.lower()
    if suffix not in DOCUMENT_SUFFIXES:
        raise ValueError(f"Unsupported document type: {suffix}")

    ensure_document_packages(install=auto_install)

    if suffix == ".pdf":
        return _read_pdf_text(path)
    return _read_docx_text(path)


def write_document_text(path: Path, content: str, *, auto_install: bool = True) -> None:
    """
    Write plain text into a PDF or DOCX file.

    :param path: Absolute output path
    :param content: Text body to persist
    :param auto_install: Attempt pip install when dependencies are missing
    :raises OptionalDependencyError: When dependencies are unavailable
    :raises ValueError: When the file type is unsupported
    """
    suffix = path.suffix.lower()
    if suffix not in DOCUMENT_SUFFIXES:
        raise ValueError(f"Unsupported document type: {suffix}")

    ensure_document_packages(install=auto_install)

    if suffix == ".pdf":
        _write_pdf_text(path, content)
        return
    _write_docx_text(path, content)


def _read_pdf_text(path: Path) -> str:
    """
    Read all pages from a PDF file.

    :param path: Absolute PDF path
    :return: Concatenated page text
    """
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(f"[Page {index}]\n{page_text.strip()}")
    return "\n\n".join(parts).strip() + ("\n" if parts else "")


def _read_docx_text(path: Path) -> str:
    """
    Read paragraph text from a DOCX file.

    :param path: Absolute DOCX path
    :return: Paragraphs joined by newlines
    """
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    body = "\n".join(paragraphs)
    return body if body.endswith("\n") else f"{body}\n"


def _write_docx_text(path: Path, content: str) -> None:
    """
    Create or overwrite a DOCX file from plain text paragraphs.

    :param path: Absolute DOCX path
    :param content: Text body
    """
    from docx import Document

    document = Document()
    for line in content.splitlines() or [""]:
        document.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def _write_pdf_text(path: Path, content: str) -> None:
    """
    Create or overwrite a simple PDF file from plain text.

    :param path: Absolute PDF path
    :param content: Text body
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    safe_content = content.encode("latin-1", errors="replace").decode("latin-1")
    if not safe_content.endswith("\n"):
        safe_content = f"{safe_content}\n"
    pdf.multi_cell(0, 8, safe_content)
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(path))
